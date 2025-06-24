# utils/report_docx.py
# -------------------------------------------------
import json                     # ←★ 추가
from docx import Document
from docx.shared import Pt

def generate_report(records: list[dict], pdf_name: str, out_path):
    """평가 결과를 Word( .docx )로 저장"""
    doc = Document()

    # 제목
    doc.add_heading("개인정보처리방침 자동 평가 결과", level=1)
    doc.add_paragraph(f"원본 PDF: {pdf_name}")

    for r in records:
        doc.add_heading(f"ID {r['eval_id']} – "
                        f"{'통과' if r['status']=='ok' else '실패'}", level=2)

        doc.add_paragraph(r["sentence"], style="Intense Quote")

        if r["status"] == "ok":
            doc.add_paragraph(
                json.dumps(r["result"], ensure_ascii=False, indent=2)
            )
        else:
            doc.add_paragraph(r.get("error", "알 수 없는 오류"))

    # 글꼴 조금만 키우기(선택)
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)

    doc.save(out_path)